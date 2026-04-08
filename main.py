from flask import Flask, render_template, request, send_file, jsonify, redirect, url_for, flash
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from flask_bcrypt import Bcrypt
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, nsmap
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
import json
import base64
import io
import uuid
import os
from models import db, User, Template

# Initialize Flask app
app = Flask(__name__, template_folder="templates")

# Configuration
app.config['SECRET_KEY'] = 'your-secret-key-change-this-in-production'  # Change this to a random string
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:////tmp/record_generator.db'  # Creates database file
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['MAX_CONTENT_LENGTH'] = 200 * 1024 * 1024

# Initialize extensions
db.init_app(app)
bcrypt = Bcrypt(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'  # Where to go if not logged in
login_manager.login_message = "Please log in to access this page."

nsmap['v'] = 'urn:schemas-microsoft-com:vml'


# User loader for Flask-Login
@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


# Create database tables
with app.app_context():
    db.create_all()
    print("Database tables created!")


# ==================== AUTHENTICATION ROUTES ====================

@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')

        # Check if user exists
        if User.query.filter_by(username=username).first():
            flash('Username already exists!', 'error')
            return redirect(url_for('register'))

        if User.query.filter_by(email=email).first():
            flash('Email already registered!', 'error')
            return redirect(url_for('register'))

        # Hash password and create user
        hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')
        new_user = User(username=username, email=email, password_hash=hashed_password)
        db.session.add(new_user)
        db.session.commit()

        flash('Account created! Please log in.', 'success')
        return redirect(url_for('login'))

    return render_template("register.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get('username')
        password = request.form.get('password')

        user = User.query.filter_by(username=username).first()

        if user and bcrypt.check_password_hash(user.password_hash, password):
            login_user(user)
            next_page = request.args.get('next')
            return redirect(next_page) if next_page else redirect(url_for('dashboard'))
        else:
            flash('Login failed. Check username and password.', 'error')

    return render_template("login.html")


@app.route("/logout")
@login_required
def logout():
    logout_user()
    flash('You have been logged out.', 'info')
    return redirect(url_for('login'))


# ==================== TEMPLATE DASHBOARD (NEW) ====================

@app.route("/dashboard")
@login_required
def dashboard():
    """Template selection page - shows user's templates"""
    # Get all templates for current user
    user_templates = Template.query.filter_by(user_id=current_user.id).order_by(Template.updated_at.desc()).all()
    return render_template("dashboard.html", templates=user_templates, username=current_user.username)


@app.route("/delete_template/<int:template_id>", methods=["POST"])
@login_required
def delete_template(template_id):
    """Delete a template"""
    template = Template.query.get_or_404(template_id)

    # Verify ownership
    if template.user_id != current_user.id:
        flash('You cannot delete this template!', 'error')
        return redirect(url_for('dashboard'))

    db.session.delete(template)
    db.session.commit()
    flash('Template deleted!', 'success')
    return redirect(url_for('dashboard'))


# ==================== YOUR EXISTING DOCUMENT CODE ====================
# (Keep all your existing functions: add_page_border, add_watermark, etc.)
# Just paste them here after the auth routes

def add_footer_to_section(section, item):
    """Add footer text to a specific section (last page only)"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.clear()

    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = footer_para.add_run(str(item.get("text", "")))
    run.font.name = item.get("font", "Calibri")
    run.font.size = Pt(int(item.get("size", 10)))
    run.font.bold = item.get("bold", False)

    rPr = run._r.get_or_add_rPr()
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')
    rPr.append(color)

    if item.get("text_enabled") and item.get("text_content"):
        text_para = footer.add_paragraph()
        text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        text_run = text_para.add_run(str(item.get("text_content", "")))
        text_run.font.name = item.get("text_font", "Calibri")
        text_run.font.size = Pt(int(item.get("text_size", 12)))

        rPr2 = text_run._r.get_or_add_rPr()
        color2 = OxmlElement('w:color')
        color2.set(qn('w:val'), '000000')
        rPr2.append(color2)


def add_page_border(section):
    xml = parse_xml(
        r'<w:pgBorders %s w:offsetFrom="page">'
        r'<w:top w:val="single" w:sz="12" w:space="24"/>'
        r'<w:left w:val="single" w:sz="12" w:space="24"/>'
        r'<w:bottom w:val="single" w:sz="12" w:space="24"/>'
        r'<w:right w:val="single" w:sz="12" w:space="24"/>'
        r'</w:pgBorders>' % nsdecls('w')
    )
    section._sectPr.append(xml)


def add_watermark(section, text):
    from docx.oxml import OxmlElement
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    pict = OxmlElement('w:pict')
    shape = OxmlElement('v:shape')
    shape.set('id', 'Watermark')
    shape.set('type', '#_x0000_t136')
    shape.set(
        'style',
        'position:absolute;'
        'width:500pt;height:120pt;'
        'rotation:315;'
        'mso-position-horizontal:center;'
        'mso-position-horizontal-relative:page;'
        'mso-position-vertical:center;'
        'mso-position-vertical-relative:page;'
    )
    shape.set('stroked', 'f')
    shape.set('fillcolor', '#d9d9d9')
    textpath = OxmlElement('v:textpath')
    textpath.set('style', "font-family:'Calibri';font-size:72pt")
    textpath.set('string', text)
    shape.append(textpath)
    pict.append(shape)
    run._r.append(pict)


def set_narrow_margins(section):
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


def add_headings(doc, headings_data):
    """Process only regular headings (no footers)"""
    print(f"Processing {len(headings_data)} headings")

    for idx, item in enumerate(headings_data):
        print(f"\n--- Heading {idx + 1} ---")

        try:
            before_lines = int(item.get("before_lines", 0))
            for _ in range(before_lines):
                doc.add_paragraph("")
        except:
            pass

        p = doc.add_paragraph()
        run = p.add_run(str(item.get("text", "")))
        run.font.name = item.get("font", "Calibri")
        run.font.size = Pt(int(item.get("size", 16)))
        run.font.bold = item.get("bold", False)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        spacing = item.get("spacing", "1")
        if spacing == "till_end":
            doc.add_page_break()
        else:
            try:
                for _ in range(int(spacing)):
                    doc.add_paragraph("")
            except:
                pass

        if item.get("text_enabled"):
            print("Adding text block...")
            try:
                before = int(item.get("text_before", 0))
                for _ in range(before):
                    doc.add_paragraph("")
            except:
                pass

            text_content = item.get("text_content", "")
            if text_content:
                tp = doc.add_paragraph()
                tr = tp.add_run(str(text_content))
                tr.font.name = item.get("text_font", "Calibri")
                tr.font.size = Pt(int(item.get("text_size", 12)))

            if item.get("text_till_end"):
                doc.add_page_break()
            else:
                try:
                    after = int(item.get("text_after", 0))
                    for _ in range(after):
                        doc.add_paragraph("")
                except:
                    pass

        if item.get("image_enabled") and item.get("image_data"):
            print("Adding image block...")
            try:
                image_data_str = item.get("image_data", "")
                if image_data_str:
                    if ',' in image_data_str:
                        image_data_str = image_data_str.split(',')[1]
                    image_data = base64.b64decode(image_data_str)
                    image_stream = io.BytesIO(image_data)
                    width_inches = float(item.get("image_width", 4))
                    height_inches = float(item.get("image_height", 3))
                    pic = doc.add_picture(image_stream, width=Inches(width_inches))
                    pic.height = Inches(height_inches)

                    try:
                        img_after = int(item.get("image_after", 0))
                        for _ in range(img_after):
                            doc.add_paragraph("")
                    except:
                        pass

                    if item.get("image_till_end"):
                        doc.add_page_break()
                    print("Image added successfully")
            except Exception as e:
                print(f"Error adding image: {str(e)}")


# ==================== MODIFIED API ROUTES ====================

@app.route("/save_template", methods=["POST"])
@login_required
def save_template():
    """Save template to database with user association"""
    try:
        data = request.json
        template_code = str(uuid.uuid4())[:8].upper()

        # If updating existing template, find it
        existing_id = data.get('template_id')
        if existing_id:
            template = Template.query.get(existing_id)
            if template and template.user_id == current_user.id:
                template.data_json = json.dumps(data)
                template.updated_at = db.func.current_timestamp()
                db.session.commit()
                return jsonify({'success': True, 'code': template.code, 'id': template.id})

        # Create new template
        new_template = Template(
            user_id=current_user.id,
            name=data.get('template_name', f'Template {template_code}'),
            code=template_code,
            data_json=json.dumps(data)
        )
        db.session.add(new_template)
        db.session.commit()

        return jsonify({
            'success': True,
            'code': template_code,
            'id': new_template.id
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({
            'success': False,
            'error': str(e)
        })


@app.route("/load_template/<code>", methods=["GET"])
@login_required
def load_template(code):
    """Load template by code (verify ownership)"""
    code = code.upper().strip()
    template = Template.query.filter_by(code=code).first()

    if not template:
        return jsonify({'success': False, 'error': 'Template not found'})

    # Check if user owns this template or if it's being shared
    #if template.user_id != current_user.id:
    #   return jsonify({'success': False, 'error': 'Access denied'})

    template_data = json.loads(template.data_json)
    template_data['template_id'] = template.id
    template_data['template_name'] = template.name

    return jsonify({
        'success': True,
        'template': template_data
    })


@app.route("/get_my_templates", methods=["GET"])
@login_required
def get_my_templates():
    """Get list of current user's templates"""
    templates = Template.query.filter_by(user_id=current_user.id).all()
    return jsonify({
        'success': True,
        'templates': [
            {
                'id': t.id,
                'code': t.code,
                'name': t.name,
                'created_at': t.created_at.strftime('%Y-%m-%d %H:%M')
            } for t in templates
        ]
    })


@app.route("/import_content", methods=["POST"])
@login_required
def import_content():
    """Import content from uploaded Word document"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'})

        file = request.files['file']
        if not file.filename.endswith('.docx'):
            return jsonify({'success': False, 'error': 'Only .docx files are supported'})

        headings_json = request.form.get('headings', '[]')
        template_headings = json.loads(headings_json)

        temp_path = f"temp_import_{uuid.uuid4()}.docx"
        file.save(temp_path)

        try:
            doc = Document(temp_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            content_map = {}

            for heading in template_headings:
                if heading.get('is_footer'):
                    continue

                heading_text = heading.get('text', '').strip()
                heading_id = heading.get('id')

                if not heading_text or heading_id is None:
                    continue

                heading_lower = heading_text.lower()
                for i, para_text in enumerate(paragraphs):
                    if para_text.lower() == heading_lower or para_text.lower().startswith(heading_lower):
                        content_lines = []
                        for j in range(i + 1, len(paragraphs)):
                            next_text = paragraphs[j]
                            is_next_heading = False
                            for other_heading in template_headings:
                                if not other_heading.get('is_footer'):
                                    other_text = other_heading.get('text', '').strip().lower()
                                    if next_text.lower() == other_text or next_text.lower().startswith(other_text):
                                        is_next_heading = True
                                        break
                            if is_next_heading:
                                break
                            content_lines.append(next_text)
                        content_map[str(heading_id)] = '\n'.join(content_lines)
                        break

            os.remove(temp_path)
            return jsonify({
                'success': True,
                'content_map': content_map,
                'matched_count': len(content_map)
            })

        except Exception as e:
            if os.path.exists(temp_path):
                os.remove(temp_path)
            raise e

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route("/import_content_text", methods=["POST"])
@login_required
def import_content_text():
    """Import content from raw text"""
    try:
        data = request.json
        raw_text = data.get('text', '')
        template_headings = data.get('headings', [])

        if not raw_text.strip():
            return jsonify({'success': False, 'error': 'No text provided'})

        if not template_headings:
            return jsonify({'success': False, 'error': 'No headings provided'})

        LINES_PER_PAGE = 45
        lines = [line.strip() for line in raw_text.split('\n') if line.strip()]
        paragraphs = []
        for line in lines:
            paragraphs.append({
                'text': line,
                'lines': 1,
                'lower': line.lower()
            })

        found_headings = []
        current_line = 0

        for para_idx, para in enumerate(paragraphs):
            for heading in template_headings:
                if heading.get('is_footer'):
                    continue

                heading_text = heading.get('text', '').strip()
                heading_id = heading.get('id')

                if not heading_text or heading_id is None:
                    continue

                if para['lower'] == heading_text.lower() or para['lower'].startswith(heading_text.lower()):
                    page_num = current_line // LINES_PER_PAGE
                    found_headings.append({
                        'id': heading_id,
                        'page': page_num,
                        'para_idx': para_idx
                    })
                    break

            current_line += para['lines']

        content_map = {}
        till_end_map = {}

        for i, current in enumerate(found_headings):
            heading_id = current['id']
            start_idx = current['para_idx']
            end_idx = found_headings[i + 1]['para_idx'] if (i + 1) < len(found_headings) else len(paragraphs)

            content_lines = []
            for j in range(start_idx + 1, end_idx):
                content_lines.append(paragraphs[j]['text'])

            content_map[str(heading_id)] = '\n'.join(content_lines)

            if (i + 1) < len(found_headings):
                next_page = found_headings[i + 1]['page']
                current_page = current['page']
                till_end_map[str(heading_id)] = (next_page > current_page)
            else:
                till_end_map[str(heading_id)] = False

        return jsonify({
            'success': True,
            'content_map': content_map,
            'till_end_map': till_end_map,
            'matched_count': len(content_map)
        })

    except Exception as e:
        import traceback
        print(traceback.format_exc())
        return jsonify({'success': False, 'error': str(e)})


@app.route("/editor")
@app.route("/")
@login_required
def editor():
    """Main editor page - now requires login"""
    return render_template("index.html")


@app.route("/editor/<template_code>")
@login_required
def editor_with_template(template_code):
    """Editor with specific template pre-loaded"""
    return render_template("index.html", load_template_code=template_code)


@app.route("/", methods=["GET", "POST"])
@login_required
def index():
    """Original index route - document generation"""
    if request.method == "POST":
        print("\n===== FORM SUBMITTED =====")

        border_enabled = request.form.get("border")
        watermark_text = request.form.get("watermark")
        headings_json = request.form.get("headings", "[]")

        try:
            headings_data = json.loads(headings_json)
        except Exception as e:
            print(f"Error parsing JSON: {e}")
            headings_data = []

        regular_headings = []
        footer_items = []

        for item in headings_data:
            if item.get("is_footer"):
                footer_items.append(item)
            else:
                regular_headings.append(item)

        doc = Document()
        section = doc.sections[0]
        layout = request.form.get("layout", "narrow")

        if layout == "narrow":
            set_narrow_margins(section)
        else:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        if border_enabled:
            add_page_border(section)

        if watermark_text:
            add_watermark(section, watermark_text)

        if footer_items and regular_headings:
            LINES_PER_PAGE = 45
            total_lines = 0

            for i, item in enumerate(regular_headings):
                total_lines += int(item.get("before_lines", 0))
                total_lines += 1

                if item.get("spacing") == "till_end":
                    total_lines = ((total_lines // LINES_PER_PAGE) + 1) * LINES_PER_PAGE
                else:
                    total_lines += int(item.get("spacing", 1))

                if item.get("text_enabled"):
                    total_lines += int(item.get("text_before", 0))
                    text_content = item.get("text_content", "")
                    if text_content:
                        wrapped_lines = sum(max(1, (len(line) // 80) + 1) for line in text_content.split('\n'))
                        total_lines += wrapped_lines

                    if item.get("text_till_end"):
                        total_lines = ((total_lines // LINES_PER_PAGE) + 1) * LINES_PER_PAGE
                    else:
                        total_lines += int(item.get("text_after", 0))

                if item.get("image_enabled") and item.get("image_data"):
                    img_height = float(item.get("image_height", 3))
                    total_lines += int(img_height * 6)

                    if item.get("image_till_end"):
                        total_lines = ((total_lines // LINES_PER_PAGE) + 1) * LINES_PER_PAGE
                    else:
                        total_lines += int(item.get("image_after", 0))

            last_heading_page = (total_lines // LINES_PER_PAGE) + 1
            print(f"Last heading will be on page: {last_heading_page}")

            if last_heading_page == 1:
                add_headings(doc, regular_headings)
                final_section = doc.sections[-1]
                final_section.footer.is_linked_to_previous = False
                if final_section.footer.paragraphs:
                    final_section.footer.paragraphs[0].clear()
                for footer_item in footer_items:
                    add_footer_to_section(final_section, footer_item)

            elif len(regular_headings) > 1:
                main_headings = regular_headings[:-1]
                final_heading = [regular_headings[-1]]

                add_headings(doc, main_headings)
                doc.add_section(WD_SECTION.NEW_PAGE)
                final_section = doc.sections[-1]

                set_narrow_margins(final_section)
                if border_enabled:
                    add_page_border(final_section)
                if watermark_text:
                    add_watermark(final_section, watermark_text)

                final_section.footer.is_linked_to_previous = False
                if final_section.footer.paragraphs:
                    final_section.footer.paragraphs[0].clear()

                add_headings(doc, final_heading)
                for footer_item in footer_items:
                    add_footer_to_section(final_section, footer_item)
            else:
                section.footer.is_linked_to_previous = False
                if section.footer.paragraphs:
                    section.footer.paragraphs[0].clear()

                add_headings(doc, regular_headings)
                for footer_item in footer_items:
                    add_footer_to_section(section, footer_item)

        elif footer_items and not regular_headings:
            section.footer.is_linked_to_previous = False
            if section.footer.paragraphs:
                section.footer.paragraphs[0].clear()
            for footer_item in footer_items:
                add_footer_to_section(section, footer_item)
        else:
            add_headings(doc, regular_headings)

        filename = "record.docx"
        doc.save(filename)
        return send_file(filename, as_attachment=True)

    return render_template("index.html")


if __name__ == "__main__":
    app.run(debug=True)
